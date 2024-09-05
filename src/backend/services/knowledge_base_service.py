import base64
from collections import defaultdict
from datetime import datetime
import io
import logging
import re
import uuid
from src.backend.kr8.vectordb.pgvector import PgVector2
from src.backend.kr8.embedder.sentence_transformer import SentenceTransformerEmbedder
from src.backend.kr8.document.reader.pdf import PDFReader
from docx import Document as DocxDocument
import pandas as pd
from src.backend.utils.analyst_selector import determine_analyst
from src.backend.kr8.document import Document as Kr8Document
from src.backend.kr8.document.reader.website import WebsiteReader
from src.backend.models.models import User
from src.backend.schemas.knowledge_base import DocumentCreate, DocumentResponse, DocumentUpdate, DocumentSearch
from sqlalchemy.orm import Session
from typing import List
import os

class KnowledgeBaseService:
    def __init__(self, db: Session, user: User):
        self.db = db
        self.user = user
        self.vector_db = PgVector2(
            collection="documents",
            db_url=os.getenv("DB_URL"),
            embedder=SentenceTransformerEmbedder(),
            user_id=user.id,
            org_id=user.organization_id
        )
        
        self.ensure_table_exists()

    def ensure_table_exists(self):
        if not self.vector_db.table_exists():
            self.vector_db.create()

    async def add_url(self, url: str) -> bool:
        scraper = WebsiteReader(max_links=2, max_depth=1)
        web_documents = scraper.read(url)
        if web_documents:
            self.vector_db.upsert(web_documents)
            return True
        return False

    def clear_knowledge_base(self) -> bool:
        return self.vector_db.clear()

    async def process_file(self, filename: str, file_content: io.BytesIO) -> DocumentResponse:
        if filename.endswith('.pdf'):
            return await self.process_pdf(filename, file_content)
        elif filename.endswith('.docx'):
            return await self.process_docx(filename, file_content)
        elif filename.endswith('.txt'):
            return await self.process_txt(filename, file_content)
        elif filename.endswith('.csv'):
            return await self.process_csv(filename, file_content.getvalue())
        elif filename.endswith(('.xlsx', '.xls')):
            content_b64 = base64.b64encode(file_content.getvalue()).decode('utf-8')
            return await self.process_excel(filename, content_b64)
        else:
            raise ValueError(f"Unsupported file type: {filename}")

    async def process_pdf(self, filename: str, file_content: io.BytesIO) -> DocumentResponse:
        reader = PDFReader()
        auto_rag_documents = reader.read(file_content, original_filename=filename)
        if not auto_rag_documents:
            raise ValueError(f"Could not read PDF: {filename}")

        self.vector_db.upsert(auto_rag_documents)

        # Return the first document as a sample
        doc = auto_rag_documents[0]
        return DocumentResponse(
            id=doc.id,
            name=doc.name,
            content=doc.content[:1000],  # Limit content for response
            meta_data=doc.meta_data,
            user_id=self.user.id,
            created_at=doc.meta_data.get('creation_date'),
            updated_at=datetime.now().isoformat()
        )

    async def process_docx(self, filename: str, file_content: io.BytesIO) -> DocumentResponse:
        doc = DocxDocument(file_content)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        content = ' '.join(full_text)
        
        kr8_doc = Kr8Document(
            content=content,
            name=filename,
            meta_data={
                "type": "docx",
                "size": len(content),
                "uploaded_at": datetime.now().isoformat()
            },
            usage={
                "access_count": 0,
                "last_accessed": None,
                "created_at": datetime.now().isoformat(),
                "updated_at": datetime.now().isoformat(),
                "relevance_scores": [],
                "token_count": None
            }
        )
        self.vector_db.upsert([kr8_doc])
        
        return DocumentResponse(
            id=kr8_doc.id,
            name=kr8_doc.name,
            content=kr8_doc.content[:1000],  # Limit content for response
            meta_data=kr8_doc.meta_data,
            user_id=self.user.id,
            created_at=kr8_doc.usage.get('created_at'),
            updated_at=kr8_doc.usage.get('updated_at')
        )

    async def process_txt(self, filename: str, file_content: io.BytesIO) -> DocumentResponse:
        content = file_content.getvalue().decode("utf-8")
        
        kr8_doc = Kr8Document(
            content=content,
            name=filename,
            meta_data={
                "type": "txt",
                "size": len(content),
                "uploaded_at": datetime.now().isoformat()
            },
            usage={
                "access_count": 0,
                "last_accessed": None,
                "created_at": datetime.now().isoformat(),
                "updated_at": datetime.now().isoformat(),
                "relevance_scores": [],
                "token_count": None
            }
        )
        self.vector_db.upsert([kr8_doc])
        
        return DocumentResponse(
            id=kr8_doc.id,
            name=kr8_doc.name,
            content=kr8_doc.content[:1000],  # Limit content for response
            meta_data=kr8_doc.meta_data,
            user_id=self.user.id,
            created_at=kr8_doc.usage.get('created_at'),
            updated_at=kr8_doc.usage.get('updated_at')
        )

    async def process_csv(self, filename: str, file_content: bytes) -> str:
        df = pd.read_csv(io.BytesIO(file_content))
        analyst_type = determine_analyst(filename, df)
        
        doc = Kr8Document(
            name=filename,
            content=df.to_csv(index=False),
            meta_data={"type": "csv", "shape": df.shape, "analyst_type": analyst_type}
        )
        self.vector_db.upsert([doc])
        
        return f"{filename} processed as {analyst_type} data"

    async def process_excel(self, filename: str, file_content_b64: str) -> str:
        file_content = base64.b64decode(file_content_b64)
        df = pd.read_excel(io.BytesIO(file_content))
        analyst_type = determine_analyst(filename, df)
        
        doc = Kr8Document(
            name=filename,
            content=df.to_csv(index=False),
            meta_data={"type": "excel", "shape": df.shape, "analyst_type": analyst_type}
        )
        self.vector_db.upsert([doc])
        
        return f"{filename} processed as {analyst_type} data"    

    def add_document(self, document: DocumentCreate) -> DocumentResponse:
        # Create a Kr8Document
        kr8_doc = Kr8Document(
            id=str(uuid.uuid4()),
            name=document.name,
            content=document.content,
            meta_data={**document.meta_data, "user_id": self.user.id} if document.meta_data else {"user_id": self.user.id}
        )
        
        # Add to vector database
        self.vector_db.insert([kr8_doc])
        
        # Create a response
        return DocumentResponse(
            id=kr8_doc.id,
            name=kr8_doc.name,
            content=kr8_doc.content,
            meta_data=kr8_doc.meta_data,
            user_id=self.user.id,
            org_id=self.user.organization_id, 
            created_at=kr8_doc.usage.get('created_at'),
            updated_at=kr8_doc.usage.get('updated_at')
        )

    def get_documents(self) -> List[DocumentResponse]:
        self.ensure_table_exists()
        document_info = self.vector_db.list_document_names()

        # Group chunks by base document name
        document_groups = defaultdict(lambda: {'chunks': 0, 'pages': set()})
        for doc_info in document_info:
            name = doc_info['name']
            # Extract base name and page number
            match = re.match(r'(.+)_page_(\d+)', name)
            if match:
                base_name, page_num = match.groups()
                document_groups[base_name]['chunks'] += doc_info['chunks']
                document_groups[base_name]['pages'].add(int(page_num))
            else:
                # If it doesn't match the pattern, treat it as a single-chunk document
                document_groups[name]['chunks'] = doc_info['chunks']

        documents = []
        for base_name, info in document_groups.items():
            meta_data = {
                "total_chunks": info['chunks'],
                "total_pages": len(info['pages']) if info['pages'] else 1
            }
            documents.append(
                DocumentResponse(
                    id=base_name,
                    name=base_name,
                    content="",  # Leave content empty
                    meta_data=meta_data,
                    user_id=self.user.id,
                    org_id=self.user.organization_id,
                    created_at=None,
                    updated_at=None,
                    chunks=info['chunks']
                )
            )

        logging.debug(f"Returning {len(documents)} documents")
        return documents

    def search_documents(self, search: DocumentSearch) -> List[DocumentResponse]:
        results = self.vector_db.search(search.query, limit=5)
        return [
            DocumentResponse(
                id=doc.id,
                name=doc.name,
                content=doc.content,  # Include the content for search results
                meta_data=doc.meta_data,
                user_id=self.user.id,
                org_id=self.user.organization_id, 
                created_at=doc.usage.get('created_at'),
                updated_at=doc.usage.get('updated_at')
            ) for doc in results if doc.meta_data.get('user_id') == self.user.id
        ]

    def delete_document(self, document_name: str) -> bool:
        # Delete all chunks associated with the document
        deleted = self.vector_db.delete_document(document_name)
        if not deleted:
            raise ValueError(f"Failed to delete document: {document_name}")
        return True

    def update_document(self, document_name: str, document_update: DocumentUpdate) -> DocumentResponse:
        document = self.vector_db.get_document_by_name(document_name)
        if not document or document.meta_data.get('user_id') != self.user.id:
            raise ValueError("Document not found")

        update_data = document_update.dict(exclude_unset=True)
        for key, value in update_data.items():
            setattr(document, key, value)

        self.vector_db.upsert([document])

        return DocumentResponse(
            id=document.id,
            name=document.name,
            content=document.content,
            meta_data=document.meta_data,
            user_id=self.user.id,
            created_at=document.usage.get('created_at'),
            updated_at=document.usage.get('updated_at')
        )