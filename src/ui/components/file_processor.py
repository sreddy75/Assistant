import base64
import io

import sqlalchemy
from kr8.document.reader.pdf import PDFReader
from kr8.document import Document
from docx import Document as DocxDocument
from datetime import datetime
from kr8.utils.log import logger
import pandas as pd
from sqlite3 import IntegrityError

def process_pdf(file, llm_os):
    reader = PDFReader()
    auto_rag_documents = reader.read(file)
    if not auto_rag_documents:
        logger.error(f"Could not read PDF: {file.name}")
        return False
    
    logger.info(f"Successfully read PDF: {file.name}. Found {len(auto_rag_documents)} documents.")
    
    try:
        llm_os.knowledge_base.load_documents(auto_rag_documents)
        logger.info("Successfully added PDF content to knowledge base.")
        return True, "Successfully added PDF content to knowledge base."
    
    except sqlalchemy.exc.ProgrammingError as e:
        if "relation" in str(e) and "does not exist" in str(e):
            # Table doesn't exist, try to create it
            llm_os.knowledge_base.vector_db.create()
            # Retry loading documents
            llm_os.knowledge_base.load_documents(auto_rag_documents)
            return True, "Created table and added PDF content to knowledge base."
        else:
            raise
    except Exception as e:
        logger.error(f"Error adding PDF content to knowledge base: {str(e)}")
        return False

def process_docx(file, llm_os):
    try:
        doc = DocxDocument(file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        content = ' '.join(full_text)
        
        doc = Document(
            content=content,
            name=file.name,
            meta_data={
                "type": "docx",
                "size": len(content),
                "uploaded_at": datetime.now().isoformat()
            },
            usage={
                "access_count": 0,
                "last_accessed": None,
                "created_at": datetime.now().isoformat(),
                "updated_at": datetime.now().isoformat(),
                "relevance_scores": [],
                "token_count": None
            }
        )
        llm_os.knowledge_base.load_documents([doc])
        logger.info(f"Processed and added DOCX file {file.name} to knowledge base")
        return True
    except Exception as e:
        logger.error(f"Error processing DOCX file {file.name}: {str(e)}")
        return False

def process_txt(file, llm_os):
    try:
        content = file.getvalue().decode("utf-8")
        
        doc = Document(
            content=content,
            name=file.name,
            meta_data={
                "type": "txt",
                "size": len(content),
                "uploaded_at": datetime.now().isoformat()
            },
            usage={
                "access_count": 0,
                "last_accessed": None,
                "created_at": datetime.now().isoformat(),
                "updated_at": datetime.now().isoformat(),
                "relevance_scores": [],
                "token_count": None
            }
        )
        llm_os.knowledge_base.load_documents([doc])
        logger.info(f"Processed and added TXT file {file.name} to knowledge base")
        return True
    except Exception as e:
        logger.error(f"Error processing TXT file {file.name}: {str(e)}")
        return False

def process_file_for_analyst(llm_os, file, file_content, analyst):
    if analyst is None:
        logger.error(f"Analyst is None when processing file: {file.name}")
        return f"Error: Unable to process {file.name} due to missing analyst"

    if not hasattr(analyst, 'get_pandas_tools') or not callable(getattr(analyst, 'get_pandas_tools')):
        logger.error(f"Analyst {analyst.name if hasattr(analyst, 'name') else 'Unknown'} has no 'get_pandas_tools' method")
        return f"Error: Unable to process {file.name} due to misconfigured analyst"

    pandas_tools = analyst.get_pandas_tools()
    if not pandas_tools:
        logger.error(f"PandasTools not available for {analyst.name if hasattr(analyst, 'name') else 'Unknown'}")
        return f"Error: Unable to process {file.name} due to missing PandasTools"
    
    try:
        if file.name.endswith('.csv'):
            df_name = pandas_tools.load_csv(file.name, file_content)
            df = pandas_tools.dataframes[df_name]

            # Convert DataFrame to Document and store in vector database
            doc = Document(
                name=file.name,
                content=df.to_csv(index=False),
                meta_data={"type": "csv", "shape": df.shape}
            )
            # Use upsert operation to handle existing documents
            try:
                llm_os.knowledge_base.vector_db.upsert([doc])
                logger.info(f"Upserted CSV file {file.name} to vector database")
            except AttributeError:
                # If upsert is not available, fall back to insert
                try:
                    llm_os.knowledge_base.vector_db.insert([doc])
                    logger.info(f"Inserted CSV file {file.name} to vector database")
                except IntegrityError:
                    logger.warning(f"Document {file.name} already exists in the database. Skipping insertion.")
            
        elif file.name.endswith(('.xlsx', '.xls')):
            df_name = pandas_tools.load_excel(file.name, file_content)
        else:
            logger.error(f"Unsupported file type: {file.name}")
            return f"Error: Unsupported file type for {file.name}"
        
        logger.info(f"Successfully loaded {file.name} as {df_name}")
        return df_name
    except Exception as e:
        logger.error(f"Error processing {file.name}: {str(e)}")
        return f"Error processing {file.name}: {str(e)}"

def determine_analyst(file, file_content):
    # Read a sample of the file to determine its content
    if file.name.endswith('.csv'):
        df = pd.read_csv(io.StringIO(file_content), nrows=5)
    elif file.name.endswith(('.xlsx', '.xls')):
        df = pd.read_excel(io.BytesIO(base64.b64decode(file_content)), nrows=5)
    else:
        return None

    # Check column names for keywords
    columns = df.columns.str.lower()
    if any(keyword in columns for keyword in ['revenue', 'financial', 'profit', 'cost']):
        return 'financial'
    else:
        return 'data'